import os
import json
import pyzipper
from io import BytesIO
from docx import Document
import streamlit as st
from dotenv import load_dotenv

# Load environment variables (for local dev)
env_path = os.path.join(os.getcwd(), '.env')

def load_config():
    """Loads configuration settings from config.json."""
    config_path = os.path.join(os.path.dirname(__file__), "config.json")
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"[Warning] Failed to read config.json: {e}")
    # Default fallback
    return {
        "subject_name": "社會學",
        "assistant_role": "專業社會學助教",
        "focus_instruction": "請著重於社會學理論、學派（如結構功能論、衝突論、符號互動論）及學者觀點之應用。",
        "materials_dir": "materials",
        "target_exam_keywords": ["高等考試三級", "地方政府公務人員", "原住民族", "身心障礙"],
        "target_subject": "社會學"
    }

def get_material_password():
    """Retrieve password from environment or Streamlit secrets."""
    # Priority 1: Streamlit Secrets (Cloud)
    try:
        if "MATERIALS_PASSWORD" in st.secrets:
            return st.secrets["MATERIALS_PASSWORD"].encode()
    except FileNotFoundError:
        pass # Secrets file not found, use env vars
    except Exception:
        pass # Other streamlit errors

    # Priority 2: Environment Variable (Local)
    return os.getenv("MATERIALS_PASSWORD", "").encode()

def load_materials():
    """
    Loads course materials.
    Priority 1: Load from unencrypted 'materials' directory (for local testing).
    Priority 2: Load from encrypted 'materials.zip' (for cloud deployment).
    """
    config = load_config()
    materials = {}
    materials_dir = config.get("materials_dir", "materials")
    zip_path = "materials.zip"
    
    # ── 1. 優先從本機的 materials 資料夾讀取（不加密） ──
    if os.path.exists(materials_dir) and os.path.isdir(materials_dir):
        try:
            docx_files = [f for f in os.listdir(materials_dir) if f.endswith('.docx')]
            if docx_files:
                for filename in docx_files:
                    week_num = filename.split('_')[0].replace('W', '').replace('-', '')
                    topic_name = f"Week {week_num}" if week_num.isdigit() else filename
                    filepath = os.path.join(materials_dir, filename)
                    doc = Document(filepath)
                    text_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    materials[topic_name] = text_content
                print(f"[Info] Loaded {len(materials)} topics from local '{materials_dir}'.")
                return materials
        except Exception as e:
            print(f"[Error] Failed to load from local directory: {e}")

    # ── 2. 若沒有資料夾，改從 materials.zip 讀取（雲端部署環境） ──
    if not os.path.exists(zip_path):
        st.error(f"❌ Error: Both '{materials_dir}/' and '{zip_path}' not found!")
        return materials

    password = get_material_password()
    if not password:
        print("[Error] No password loaded! Check .env or secrets.")
    else:
        print("[Info] ZIP Password loaded.")

    try:
        with pyzipper.AESZipFile(zip_path, 'r') as zf:
            if password:
                zf.setpassword(password)
            docx_files = [name for name in zf.namelist() if name.endswith('.docx')]
            for filename in docx_files:
                week_num = filename.split('_')[0].replace('W', '').replace('-', '')
                topic_name = f"Week {week_num}" if week_num.isdigit() else filename
                file_data = zf.read(filename)
                doc = Document(BytesIO(file_data))
                text_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                materials[topic_name] = text_content
        print(f"[Info] Loaded {len(materials)} topics from '{zip_path}'.")
        return materials
    except RuntimeError as e:
        print(f"[Error] Decryption failed! Wrong password or corrupted file: {e}")
        return materials
    except Exception as e:
        print(f"[Error] Loading ZIP materials failed: {e}")
        return materials


if __name__ == "__main__":
    # Test run
    data = load_materials()
    print(f"Total topics loaded: {len(data)}")
    for key in data:
        print(f"- {key}: {len(data[key])} chars")
